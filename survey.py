from flask import Flask, request, jsonify, send_file, session, render_template
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import os
import uuid
import requests
from docx import Document
from docx.shared import Inches
from flask_session import Session
import matplotlib.font_manager as fm
from docx.oxml.ns import qn
from sentence_transformers import SentenceTransformer, util
from difflib import get_close_matches
from collections import Counter
from werkzeug.middleware.proxy_fix import ProxyFix
import re

# Setup
app = Flask(__name__, static_folder=".", static_url_path="")
app.secret_key = str(uuid.uuid4())
app.config['SESSION_TYPE'] = 'filesystem'
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
Session(app)

local_model = SentenceTransformer("local_miniLM_model")

font_paths = [
    'C:/Windows/Fonts/Nirmala.ttf',
    'C:/Windows/Fonts/nirmala.ttf',
    'C:/Windows/Fonts/nirmalui.ttf',
    'C:/Windows/Fonts/Nirmala UI.ttf',
    'C:/Windows/Fonts/Mangal.ttf'
]
hindi_font_path = next((p for p in font_paths if os.path.exists(p)), None)
hindi_font = fm.FontProperties(fname=hindi_font_path) if hindi_font_path else fm.FontProperties(family='Nirmala UI')

df = None
summary_text = ""

# Ollama integration
def ask_ollama(prompt):
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": "mistral", "prompt": prompt, "stream": False}
        )
        return response.json()["response"].strip()
    except Exception as e:
        return f"❌ Ollama error: {e}"

def normalize_multiselect(val):
    val = str(val).strip()
    val = re.sub(r'(?<=\d)\s*,\s*(?=\d)', '', val)
    val = val.replace("fasilan", "fislan")
    if re.search(r'[\uFFFD\u25A1\uFFFD�□]', val):
        return ""
    parts = [p.strip() for p in val.split(",") if p.strip()]
    parts.sort()
    return ", ".join(parts)

def format_display_label(val):
    return val.replace("fasilan", "fislan")

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    global df, summary_text
    uploaded_file = request.files['file']
    file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.filename)
    uploaded_file.save(file_path)

    try:
        df_raw = pd.read_excel(file_path, header=1, engine="openpyxl").iloc[2:].reset_index(drop=True)
        df_raw = df_raw.drop(index=0, errors='ignore')
        df = df_raw.dropna(how='all').reset_index(drop=True)
        df.columns = df.columns.astype(str).str.strip().str.replace("\n", " ").str.replace(r"\s+", " ", regex=True)

        # 🔑 Track current topic based on filename (without extension)
        topic_name = os.path.splitext(uploaded_file.filename)[0].lower()
        session["current_topic"] = topic_name
        if "all_custom_reports" not in session:
            session["all_custom_reports"] = {}
        session["all_custom_reports"][topic_name] = []

        summary_text = ""
        for col in df.columns:
            dtype = df[col].dtype
            top_vals = df[col].dropna().astype(str).value_counts().head(3).to_dict()
            sample_vals = df[col].dropna().astype(str).sample(min(2, len(df[col].dropna()))).tolist()
            summary_text += f"Column: {col}\nType: {dtype}\nTop Responses: {top_vals}\nSamples: {sample_vals}\n\n"

        return jsonify({"status": "success", "columns": list(df.columns)})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

app.wsgi_app = ProxyFix(app.wsgi_app)
@app.route("/query", methods=["POST"])
def query():
    user_question = request.json.get("question", "")
    include_in_report = request.json.get("include", False)

    if df is None:
        return jsonify({"answer": "⚠️ Please upload a survey file first.", "chart": None})

    try:
        question_emb = local_model.encode(user_question, convert_to_tensor=True)
        scores = [(col, util.pytorch_cos_sim(question_emb, local_model.encode(col, convert_to_tensor=True)).item()) for col in df.columns]
        best_col, best_score = max(scores, key=lambda x: x[1]) if scores else (None, 0)
        match_type = "🔍 Matched using semantic embeddings"

        if best_score < 0.5 or best_col is None or df[best_col].dropna().nunique() < 2:
            col_names = ", ".join([f'"{c}"' for c in df.columns])
            df_sample = df.head(5).to_dict(orient='records')
            prompt = f"""You are a skilled data analyst. Your job is to match a user question to a column in the dataset.

Here is a preview of the dataset:
{df_sample}

These are the exact column names available:
{col_names}

User's question:
"{user_question}"

⚠️ Very important: Return only the exact column name that best answers the user's question.
❌ Do NOT rephrase the question. Do NOT explain. Do NOT add punctuation or quotation marks.
✅ Just return the column name exactly as written in the list above."""
            ollama_response = ask_ollama(prompt)
            candidate_col = ollama_response.splitlines()[0].strip()

            valid_col = get_close_matches(candidate_col, df.columns, n=1, cutoff=0.5)
            if not valid_col:
                candidate_emb = local_model.encode(candidate_col, convert_to_tensor=True)
                scores = [(col, util.pytorch_cos_sim(candidate_emb, local_model.encode(col, convert_to_tensor=True)).item()) for col in df.columns]
                best_col, best_score = max(scores, key=lambda x: x[1])
                if best_score < 0.4:
                    return jsonify({
                        "answer": f"🤖 Ollama could not identify a valid column.\n\nResponse:\n{ollama_response}",
                        "chart": None
                    })
                match_type = "🤖 Fallback semantic match from Ollama response"
            else:
                best_col = valid_col[0]
                match_type = "🤖 Matched using Mistral (Ollama)"

        df[best_col] = df[best_col].astype(str).str.strip().apply(normalize_multiselect)
        value_counts = df[best_col][df[best_col] != ""].value_counts()
        total_responses = len(df[best_col].dropna())

        chart_path = f"static/{str(uuid.uuid4())}.png"
        summary_sentences = []

        plt.figure(figsize=(12, 8))
        if len(value_counts) <= 6:
            plt.pie(
                value_counts,
                labels=[format_display_label(label) for label in value_counts.index],
                autopct='%1.1f%%',
                startangle=140,
                textprops={'fontproperties': hindi_font, 'fontsize': 14}
            )
            plt.title(f"Responses to: {best_col}", fontproperties=hindi_font, fontsize=18)
            plt.axis('equal')
        else:
            # Remove labels with unreadable box characters
            clean_value_counts = value_counts[~value_counts.index.str.contains(r'[\uFFFD\u25A1�□]', na=False)]
            top_10 = clean_value_counts.head(10)
            sns.set(style="whitegrid")
            ax = sns.barplot(
                y=top_10.index.astype(str),
                x=top_10.values,
                orient='h'
            )
            ax.set_ylabel("")
            labels = [format_display_label(label.get_text()) for label in ax.get_yticklabels()]
            for label_obj, new_text in zip(ax.get_yticklabels(), labels):
                label_obj.set_text(new_text)
                label_obj.set_fontproperties(hindi_font)
                label_obj.set_fontsize(12)
            ax.set_xlabel("Count", fontproperties=hindi_font, fontsize=14)
            ax.set_title(f"Top 10 Responses to: {best_col}", fontproperties=hindi_font, fontsize=16)
            plt.tight_layout()

        plt.savefig(chart_path, bbox_inches='tight')
        plt.close()

        for k, v in value_counts.head(5).items():
            summary_sentences.append(f"{v} people selected '{format_display_label(k)}'")
        if len(value_counts) > 5:
            summary_sentences.append(f"...and {len(value_counts) - 5} more unique responses")

        answer = f"""{match_type}

The question: '{user_question}' relates most to column: '{best_col}'.
Total responses analyzed: {total_responses}

Top Responses:
{chr(10).join(summary_sentences)}"""

        if include_in_report:
            report_item = {
                "question": user_question,
                "answer": answer,
                "chart": chart_path
            }

            topic = session.get("current_topic", "default")
            if "all_custom_reports" not in session:
                session["all_custom_reports"] = {}
            if topic not in session["all_custom_reports"]:
                session["all_custom_reports"][topic] = []

            reports = session["all_custom_reports"][topic]
            if not any(r["question"] == user_question for r in reports):
                reports.append(report_item)
                session["all_custom_reports"][topic] = reports
                session.modified = True
        return jsonify({
            "answer": answer,
            "chart": f"/{chart_path}" if chart_path else None
        })

    except Exception as e:
        return jsonify({"answer": f"❌ Error: {str(e)}", "chart": None})

@app.before_request
def make_session_permanent():
    session.permanent = True

@app.route("/generate_summary", methods=["GET"])
def generate_summary():
    global df
    if df is None:
        return jsonify({"summary": "⚠️ No data uploaded."})
    df_sample = df.head(10).to_dict(orient='records')
    prompt = f"You are a survey data analyst. Here's a sample of the data:\n{df_sample}\n\nSummarize the key findings, interesting patterns, or questions to investigate."
    summary = ask_ollama(prompt)
    return jsonify({"summary": summary})

@app.route("/download_custom_report")
def download_custom_report():
    doc = Document()
    doc.add_heading("Custom Survey Report", 0)

    topic = session.get("current_topic", "default")
    reports_by_topic = session.get("all_custom_reports", {})
    report_items = reports_by_topic.get(topic, [])

    if not report_items:
        return "⚠️ No questions selected for custom report.", 400

    for entry in report_items:
        question = entry.get("question", "Unknown Question")
        answer = entry.get("answer", "No analysis available.")
        chart_path = entry.get("chart")

        doc.add_heading(question, level=2)
        para = doc.add_paragraph(answer)
        run = para.runs[0]
        run.font.name = 'Nirmala UI'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')

        if chart_path and os.path.exists(chart_path) and chart_path.endswith((".png", ".jpg", ".jpeg")):
            try:
                doc.add_picture(chart_path, width=Inches(5.5))
            except Exception as pic_err:
                doc.add_paragraph(f"⚠️ Could not load chart: {pic_err}")

        doc.add_paragraph("\n")

    file_path = f"static/custom_report_{topic}.docx"
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

@app.route("/download_complete_report")
def download_complete_report():
    if df is None:
        return "No survey data uploaded.", 400

    doc = Document()

    for col in df.columns:
        if str(col).strip().lower() in ["question", ""]:
            continue

        doc.add_heading(col, level=2)
        responses = df[col].dropna().astype(str).str.strip()

        def normalize_multiselect(val):
            val = str(val).strip().lower()

            # 🧼 Fix scrambled number ranges like "000, 000 to 8, rs. 6"
            val = re.sub(r'(?<=\d)\s*,\s*(?=\d)', '', val)  # remove commas inside numbers
            val = re.sub(r'(rs\.\s*\d{1,2})(\d{3})', lambda m: m.group(1) + "," + m.group(2), val)
            val = re.sub(r'(\d{1,2})(\d{3})', lambda m: m.group(1) + "," + m.group(2), val)  # 6000 → 6,000

            # 🔁 Normalize multiselect: "walker, stick" = "stick, walker"
            parts = [p.strip() for p in val.split(",") if p.strip()]
            parts.sort()
            return ", ".join(parts)

        responses = responses.apply(normalize_multiselect)

        value_counts = responses.value_counts()
        total_responses = len(responses)
        top_summary = value_counts.head(5)

        summary = f"Total Responses: {total_responses}\n"
        for val, count in top_summary.items():
            summary += f"- {count} people selected: '{val}'\n"

        para = doc.add_paragraph(summary)
        run = para.runs[0]
        run.font.name = 'Nirmala UI'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')

        if not value_counts.empty:
            chart_path = f"static/{str(uuid.uuid4())}.png"
            plt.figure(figsize=(12, 8))
            if len(value_counts) <= 6:
                plt.pie(
                    value_counts,
                    labels=value_counts.index.astype(str),
                    autopct='%1.1f%%',
                    startangle=140,
                    textprops={'fontproperties': hindi_font, 'fontsize': 12}
                )
                plt.title(f"Responses to: {col}", fontproperties=hindi_font, fontsize=16)
                plt.axis('equal')
            else:
                top_10 = value_counts.head(10)
                sns.set(style="whitegrid")
                ax = sns.barplot(
                    y=top_10.index.astype(str),
                    x=top_10.values,
                    orient='h'
                )
                ax.set_yticklabels(ax.get_yticklabels(), fontproperties=hindi_font, fontsize=12)
                ax.set_xlabel("Count", fontproperties=hindi_font, fontsize=14)
                ax.set_title(f"Top 10 Responses to: {col}", fontproperties=hindi_font, fontsize=16)
                plt.tight_layout()

            plt.savefig(chart_path, bbox_inches='tight')
            plt.close()
            doc.add_picture(chart_path, width=Inches(5.5))

    file_path = "static/complete_survey_report.docx"
    doc.save(file_path)
    return send_file(file_path, as_attachment=True)

if __name__ == "__main__":
    if not os.path.exists("static"):
        os.makedirs("static")
    app.run(debug=True, threaded=True)
