from flask import Flask, request, jsonify, send_file, session, render_template
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import os
import uuid
import difflib
from docx import Document
from docx.shared import Inches
from flask_session import Session
import matplotlib.font_manager as fm
from docx.oxml.ns import qn
from sentence_transformers import SentenceTransformer, util

app = Flask(__name__, static_folder=".", static_url_path="")
app.secret_key = str(uuid.uuid4())
app.config['SESSION_TYPE'] = 'filesystem'
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
Session(app)

# Local NLP Model Setup
local_model = SentenceTransformer("local_miniLM_model")

# Font setup
font_paths = [
    'C:/Windows/Fonts/Nirmala.ttf',
    'C:/Windows/Fonts/nirmala.ttf',
    'C:/Windows/Fonts/nirmalui.ttf',
    'C:/Windows/Fonts/Nirmala UI.ttf',
    'C:/Windows/Fonts/Mangal.ttf'
]
hindi_font_path = next((p for p in font_paths if os.path.exists(p)), None)
hindi_font = fm.FontProperties(fname=hindi_font_path) if hindi_font_path else fm.FontProperties(family='Nirmala UI')

# Globals
df = None
summary_text = ""

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    global df, summary_text
    uploaded_file = request.files['file']
    file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.filename)
    uploaded_file.save(file_path)

    try:
        df = pd.read_excel(file_path, header=2, engine="openpyxl")
        df = df.dropna(how='all').reset_index(drop=True)
        df.columns = df.columns.astype(str).str.strip().str.replace("\n", " ").str.replace(r"\s+", " ", regex=True)

        summary_text = ""
        for col in df.columns:
            dtype = df[col].dtype
            top_vals = df[col].dropna().astype(str).value_counts().head(3).to_dict()
            sample_vals = df[col].dropna().astype(str).sample(min(2, len(df[col].dropna()))).tolist()
            summary_text += f"Column: {col}\nType: {dtype}\nTop Responses: {top_vals}\nSamples: {sample_vals}\n\n"

        return jsonify({"status": "success", "columns": list(df.columns)})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/query", methods=["POST"])
def query():
    user_question = request.json.get("question", "")
    include_in_report = request.json.get("include", False)

    if df is None:
        return jsonify({"answer": "⚠️ Please upload a survey file first.", "chart": None})

    try:
        question_emb = local_model.encode(user_question, convert_to_tensor=True)
        scores = []
        for col in df.columns:
            col_emb = local_model.encode(col, convert_to_tensor=True)
            score = util.pytorch_cos_sim(question_emb, col_emb).item()
            scores.append((col, score))
        best_col = max(scores, key=lambda x: x[1])[0] if scores else None

        if not best_col:
            return jsonify({"answer": "❌ No relevant column found for this question.", "chart": None})

        value_counts = df[best_col].value_counts()
        total_responses = len(df[best_col].dropna())
        chart_path = None
        description = ""

        if not value_counts.empty:
            chart_path = f"static/{str(uuid.uuid4())}.png"
            plt.figure(figsize=(10, 5))
            if len(value_counts) <= 6:
                plt.pie(value_counts, labels=value_counts.index, autopct="%1.1f%%", startangle=140,
                        textprops={'fontproperties': hindi_font})
                plt.axis('equal')
            else:
                sns.barplot(x=value_counts.index.astype(str), y=value_counts.values)
                plt.xticks(rotation=30, ha='right', fontsize=9, fontproperties=hindi_font)
                plt.ylabel("Count", fontproperties=hindi_font)
            plt.title(best_col, fontproperties=hindi_font)
            plt.tight_layout()
            plt.savefig(chart_path)
            plt.close()

            description += "\n".join([f"{k}: {v} responses" for k, v in value_counts.items()])

        answer = f"The question '{user_question}' relates most to column: {best_col}. Total responses analyzed: {total_responses}.\n\nResponse Breakdown:\n{description}"

        if include_in_report:
            if 'custom_report' not in session:
                session['custom_report'] = []
            session['custom_report'].append({
                "question": user_question,
                "answer": answer,
                "chart": chart_path
            })
            session.modified = True

        return jsonify({"answer": answer, "chart": chart_path})

    except Exception as e:
        return jsonify({"answer": f"❌ Error: {str(e)}", "chart": None})

@app.route("/download_custom_report")
def download_custom_report():
    doc = Document()
    doc.add_heading("Custom Survey Report", 0)
    for entry in session.get("custom_report", []):
        doc.add_heading(entry["question"], level=2)
        para = doc.add_paragraph(entry["answer"])
        run = para.runs[0]
        run.font.name = 'Nirmala UI'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')
        if entry["chart"] and os.path.exists(entry["chart"]):
            doc.add_picture(entry["chart"], width=Inches(5.5))

    file_path = "static/custom_report.docx"
    doc.save(file_path)
    return send_file(file_path, as_attachment=True)

if __name__ == "__main__":
    if not os.path.exists("static"):
        os.makedirs("static")
    app.run(debug=True, threaded=True)
